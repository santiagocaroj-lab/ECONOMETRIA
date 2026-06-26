# Importación de librerías necesarias
import streamlit as st # Framework para crear la aplicación web de manera interactiva
import pandas as pd # Para el manejo, indexación y transformación de datos (DataFrames)
import numpy as np # Para operaciones matriciales y manejo especializado de nulos (NaN)
import plotly.express as px # Para la visualización interactiva de gráficos dinámicos
import seaborn as sns # Para la generación de la matriz de calor (Heatmap)
import matplotlib.pyplot as plt # Soporte estructural para gráficos estáticos de Seaborn
import statsmodels.api as sm # Para agregar la constante econométrica al set de regresores
from linearmodels.panel import PanelOLS # Para de panel con doble efecto fijo
import docx # Para compilar y exportar reportes automatizados en formato Word
from docx.shared import Inches # Ajuste de márgenes y dimensiones en el reporte exportable
import io # Para la gestión eficiente de buffers de datos binarios en la memoria virtual

# -----------------------------------------------------------------------------
# CONFIGURACIÓN DE LA INTERFAZ DE USUARIO Y BARRA LATERAL
# -----------------------------------------------------------------------------
st.set_page_config(page_title="Análisis Econométrico de Panel", layout="wide")
st.title("📊 Análisis Econométrico: Libertad de Prensa, Corrupción e IED")
st.markdown("""
Esta aplicación permite transformar una base de datos, explorar descriptivamente las interacciones, 
identificar casos atípicos y estimar modelos de panel bidimensionales e invertidos. Utiliza el menú lateral para ajustar los parámetros globales.
""")

st.sidebar.header("1. Carga de Archivos")
uploaded_data = st.sidebar.file_uploader("Sube tu base de datos (Excel o CSV)", type=["xlsx", "csv"])
uploaded_doc = st.sidebar.file_uploader("Sube el documento Word metodológico (Opcional)", type=["docx"])

if uploaded_doc is not None:
    st.sidebar.success("Documento Word cargado correctamente.")
    with st.sidebar.expander("📄 Ver Documento Metodológico"):
        doc = docx.Document(uploaded_doc)
        for para in doc.paragraphs:
            st.write(para.text)

st.sidebar.markdown("---")
st.sidebar.header("2. Configuración Global")
# TOGGLE MAESTRO: Controla si toda la app usa (t-1) o (t) para dispersión y modelos
aplicar_rezago_global = st.sidebar.checkbox(
    "🔄 Aplicar Rezago Temporal (t-1)", 
    value=True, 
    help="Si está activo, evalúa cómo el año anterior impacta al actual (Modelos Dinámicos). Si se desactiva, evalúa relaciones en el mismo año (Modelos Contemporáneos Estáticos)."
)

if aplicar_rezago_global:
    st.sidebar.info("Modo Dinámico Activado: Las variables explicativas tendrán un rezago de 1 año (t-1).")
else:
    st.sidebar.warning("Modo Estático Activado: Todas las relaciones se evalúan en el mismo año (t).")

# -----------------------------------------------------------------------------
# MOTOR DE PROCESAMIENTO ECONOMÉTRICO
# -----------------------------------------------------------------------------
if uploaded_data is not None:
    try:
        if uploaded_data.name.endswith('.csv'):
            df_wide = pd.read_csv(uploaded_data)
        else:
            df_wide = pd.read_excel(uploaded_data)
    except Exception as e:
        st.error(f"Error crítico al leer el archivo de entrada: {e}")
        st.stop()

    # Definición de una lista explícita de stubnames para máxima robustez
    stubnames_lista = ['CPI_SCORE', 'CPI_RANK', 'CPI_RANKING', 'CORRUPCION', 'RSF_SCORE', 'RSF_RANK', 'RSF_RANKING', 'IED_PIB']
    
    # Reestructuración limpia de formato Ancho (Wide) a Largo (Long)
    df_long = pd.wide_to_long(
        df_wide, 
        stubnames=stubnames_lista,
        i='País',
        j='Año',
        sep='_',
        suffix=r'\d+'
    ).reset_index()               

    # Ordenamiento jerárquico indispensable para el correcto cálculo de rezagos temporales
    df_long = df_long.sort_values(by=['País', 'Año'])

    # -------------------------------------------------------------------------
    # GENERACIÓN DE VARIABLES REZAGADAS (LAGGED VARIABLES)
    # -------------------------------------------------------------------------
    df_long['RSF_L1'] = df_long.groupby('País')['RSF_SCORE'].shift(1)
    df_long['CORRUPCION_L1'] = df_long.groupby('País')['CORRUPCION'].shift(1)
    df_long['CPI_L1'] = df_long.groupby('País')['CPI_SCORE'].shift(1)
    
    if 'IED_PIB' in df_long.columns:
        df_long['IED_L1'] = df_long.groupby('País')['IED_PIB'].shift(1)
    else:
        df_long['IED_L1'] = np.nan

    # Asignación dinámica de variables explicativas basadas en el Toggle Maestro
    var_rsf_explicativa = 'RSF_L1' if aplicar_rezago_global else 'RSF_SCORE'
    var_corr_explicativa = 'CORRUPCION_L1' if aplicar_rezago_global else 'CORRUPCION'
    var_cpi_explicativa = 'CPI_L1' if aplicar_rezago_global else 'CPI_SCORE'
    sufijo_tiempo = "(t-1)" if aplicar_rezago_global else "(t)"

    # Despliegue modular a través de pestañas funcionales
    tab1, tab2, tab3, tab4 = st.tabs(["Estadísticas Descriptivas", "Gráficos Interactivos", "Modelos Econométricos", "Exportar Resultados"])

    # =========================================================================
    # PESTAÑA 1: ESTADÍSTICAS DESCRIPTIVAS
    # =========================================================================
    with tab1:
        st.subheader("Estadísticas Descriptivas del Panel")
        desc_stats = df_long.describe().T[['mean', '50%', 'std', 'min', 'max']]
        desc_stats.columns = ['Media', 'Mediana', 'Desv. Estándar', 'Mínimo', 'Máximo']
        st.dataframe(desc_stats)

        st.subheader("Matriz de Correlaciones")
        df_numeric = df_long.select_dtypes(include=[np.number])
        corr_matrix = df_numeric.corr()
        
        fig_corr, ax_corr = plt.subplots(figsize=(10, 6))
        sns.heatmap(corr_matrix, annot=True, cmap="coolwarm", fmt=".2f", linewidths=0.5, ax=ax_corr)
        st.pyplot(fig_corr)

    # =========================================================================
    # PESTAÑA 2: GRÁFICOS INTERACTIVOS, CASOS ATÍPICOS Y DISPERSIÓN
    # =========================================================================
    with tab2:
        st.subheader("1. Evolución Temporal de Indicadores (Personalizable)")
        
        col_ctrl1, col_ctrl2 = st.columns(2)
        with col_ctrl1:
            nivel_analisis = st.radio("Seleccione el nivel de análisis:", ["Promedio Regional", "País Específico"])
            if nivel_analisis == "País Específico":
                pais_seleccionado = st.selectbox("Seleccione un país:", df_long['País'].unique())
        
        with col_ctrl2:
            columnas_grafico = [col for col in ['CPI_SCORE', 'RSF_SCORE', 'CORRUPCION', 'IED_PIB'] if col in df_long.columns]
            variables_seleccionadas = st.multiselect(
                "Seleccione las variables a graficar:", 
                options=columnas_grafico, 
                default=columnas_grafico
            )
            
        if not variables_seleccionadas:
            st.warning("Por favor, seleccione al menos una variable para visualizar.")
        else:
            if nivel_analisis == "Promedio Regional":
                df_grafico = df_long.groupby('Año')[variables_seleccionadas].mean().reset_index()
                titulo = 'Evolución Promedio en la Región'
            else:
                df_grafico = df_long[df_long['País'] == pais_seleccionado]
                titulo = f'Trayectoria Temporal en {pais_seleccionado}'
                
            fig_line = px.line(
                df_grafico, x='Año', y=variables_seleccionadas, title=titulo, markers=True,
                color_discrete_sequence=['#1f77b4', '#ff7f0e', '#d62728', '#2ca02c'] 
            )
            
            fig_line.update_layout(
                template='simple_white',
                title={'x': 0.5, 'xanchor': 'center', 'font': {'size': 18, 'family': 'Arial', 'color': 'black'}},
                xaxis_title="Año", yaxis_title="Índice / Porcentaje", legend_title_text="Variables:",
                font=dict(family="Arial", size=13, color="black"), hovermode="x unified",
                margin=dict(l=50, r=30, t=60, b=50), legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1)
            )
            fig_line.update_traces(line=dict(width=2.5), marker=dict(size=8))
            st.plotly_chart(fig_line, use_container_width=True)

        st.markdown("---")
        st.subheader("2. Casos Atípicos y Análisis de Variación (Ganadores y Perdedores)")
        st.write("Visualiza rápidamente qué países registraron los cambios más extremos (positivos o negativos) comparando su primer año de registro con su último año.")
        
        var_atipica = st.selectbox("Seleccione la variable para analizar casos atípicos:", columnas_grafico)
        
        df_agrupado = df_long.dropna(subset=[var_atipica]).sort_values(by=['País', 'Año'])
        primeros_registros = df_agrupado.groupby('País').first().reset_index()
        ultimos_registros = df_agrupado.groupby('País').last().reset_index()
        
        df_variaciones = pd.DataFrame({
            'País': primeros_registros['País'],
            'Año Inicial': primeros_registros['Año'],
            'Valor Inicial': primeros_registros[var_atipica],
            'Año Final': ultimos_registros['Año'],
            'Valor Final': ultimos_registros[var_atipica],
            'Variación Neta': ultimos_registros[var_atipica] - primeros_registros[var_atipica]
        }).sort_values(by='Variación Neta', ascending=False)
        
        fig_var = px.bar(
            df_variaciones, x='País', y='Variación Neta', title=f"Variación Total de {var_atipica} en el Período Estudiado",
            color='Variación Neta', color_continuous_scale=px.colors.diverging.RdYlGn, text_auto='.2f'
        )
        fig_var.update_layout(template='simple_white', font=dict(family="Arial", color="black"), xaxis_title="", yaxis_title="Cambio Neto", coloraxis_showscale=False)
        st.plotly_chart(fig_var, use_container_width=True)
        
        col_ganadores, col_perdedores = st.columns(2)
        with col_ganadores:
            st.success(f"🏆 Top 3: Mayores Incrementos en {var_atipica}")
            st.dataframe(df_variaciones.head(3)[['País', 'Valor Inicial', 'Valor Final', 'Variación Neta']].style.format({"Valor Inicial": "{:.2f}", "Valor Final": "{:.2f}", "Variación Neta": "{:+.2f}"}))
        with col_perdedores:
            st.error(f"🚨 Top 3: Caídas Más Extremas en {var_atipica}")
            st.dataframe(df_variaciones.tail(3)[['País', 'Valor Inicial', 'Valor Final', 'Variación Neta']].style.format({"Valor Inicial": "{:.2f}", "Valor Final": "{:.2f}", "Variación Neta": "{:+.2f}"}))

        st.markdown("---")
        st.subheader(f"3. Diagramas de Dispersión Dinámicos (Efecto Global: {sufijo_tiempo})")
        
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            fig_scatter1 = px.scatter(df_long, x=var_rsf_explicativa, y='CORRUPCION', hover_data=['País', 'Año'], trendline='ols', title=f"RSF {sufijo_tiempo} vs Corrupción Real (t)")
            st.plotly_chart(fig_scatter1, use_container_width=True)
        with col2:
            fig_scatter2 = px.scatter(df_long, x=var_rsf_explicativa, y='CPI_SCORE', hover_data=['País', 'Año'], trendline='ols', title=f"RSF {sufijo_tiempo} vs CPI Score (t)")
            st.plotly_chart(fig_scatter2, use_container_width=True)
        with col3:
            if 'IED_PIB' in df_long.columns:
                fig_scatter3 = px.scatter(df_long, x=var_corr_explicativa, y='IED_PIB', hover_data=['País', 'Año'], trendline='ols', title=f"Corrupción {sufijo_tiempo} vs IED (t)")
                st.plotly_chart(fig_scatter3, use_container_width=True)
        with col4:
            if 'IED_PIB' in df_long.columns:
                fig_scatter4 = px.scatter(df_long, x=var_cpi_explicativa, y='IED_PIB', hover_data=['País', 'Año'], trendline='ols', title=f"CPI Score {sufijo_tiempo} vs IED (t)")
                st.plotly_chart(fig_scatter4, use_container_width=True)

    # =========================================================================
    # PESTAÑA 3: ESTIMACIÓN DE MODELOS MATRICIALES (INVERTIDOS Y DUALES)
    # =========================================================================
    with tab3:
        tipo_modelo_texto = "DINÁMICOS (Efectos Fijos Bidimensionales con Rezagos t-1)" if aplicar_rezago_global else "ESTÁTICOS CONTEMPORÁNEOS (Efectos Fijos Bidimensionales t)"
        st.subheader(f"Estimación Matriz Econométrica: Modelos {tipo_modelo_texto}")
        st.write("Esta sección ejecuta las regresiones controlando por efectos fijos de país y año de manera simultánea. Permite comparar la robustez de usar la Corrupción Directa vs. el CPI Score.")
            
        df_panel = df_long.set_index(['País', 'Año'])

        def interpretar_resultado_avanzado(tipo_modelo, variable_foco, coef, p_value, usar_rezago):
            significativo = p_value < 0.05
            t_anterior = "en el año anterior impacta" if usar_rezago else "se asocia contemporáneamente con"
            
            texto = f"**Interpretación Institucional ({variable_foco}):**\n\n"
            
            if tipo_modelo == "RSF_to_CORR":
                if significativo:
                    if coef < 0:
                        texto += f"El coeficiente es **negativo** ({coef:.4f}) y significativo (p < 0.05). Indica que mayor libertad de prensa {t_anterior} una **reducción de la corrupción real**, validando la teoría de fiscalización social de los medios."
                    else:
                        texto += f"El coeficiente es **positivo** ({coef:.4f}) y significativo. Alerta sobre un 'efecto destape': la apertura de prensa visibiliza y expone mediáticamente más casos."
                else:
                    texto += "No se observa un impacto estadísticamente significativo bajo esta especificación de panel de efectos fijos."
                    
            elif tipo_modelo == "RSF_to_CPI":
                if significativo:
                    if coef > 0:
                        texto += f"El coeficiente es **positivo** ({coef:.4f}) y significativo (p < 0.05). Mayor libertad de prensa {t_anterior} un **incremento en la transparencia (CPI Score)**. Es perfectamente consistente con el modelo de corrupción real (signos opuestos debido a la construcción del índice)."
                    else:
                        texto += f"El coeficiente es **negativo** ({coef:.4f}) y significativo. El índice de transparencia se deteriora ante mayor libertad."
                else:
                    texto += "El impacto de la libertad de prensa sobre el CPI Score no es estadísticamente significativo."
                    
            elif tipo_modelo == "CORR_to_RSF":
                if significativo:
                    if coef < 0:
                        texto += f"El coeficiente es **negativo** ({coef:.4f}) y significativo. Evidencia un canal de **captura institucional**: mayores niveles de corrupción real {t_anterior} un menoscabo posterior de las garantías para el ejercicio periodístico."
                    else:
                        texto += f"El coeficiente es **positivo** ({coef:.4f}). Estructuras corruptas se asocian con mayor libertad de prensa."
                else:
                    texto += "La hipótesis de retroalimentación inversa desde la corrupción real hacia la prensa no es significativa."
                    
            elif tipo_modelo == "CPI_to_RSF":
                if significativo:
                    if coef > 0:
                        texto += f"El coeficiente es **positivo** ({coef:.4f}) y significativo. Demuestra que entornos con mayor transparencia y seguridad jurídica (CPI Score elevado) {t_anterior} un **blindaje y fortalecimiento de la libertad de prensa**."
                    else:
                        texto += f"El coeficiente es **negativo** ({coef:.4f}). Mayor transparencia reduce el score de prensa."
                else:
                    texto += "La retroalimentación desde el CPI Score hacia la prensa no muestra significancia estadística."
                    
            elif tipo_modelo == "CORR_to_IED":
                if significativo:
                    if coef < 0:
                        texto += f"El coeficiente es **negativo** ({coef:.4f}) y significativo. Valida la tesis del **riesgo soberano e incertidumbre**: la corrupción real actúa como un impuesto implícito que ahuyenta los flujos de IED."
                    else:
                        texto += f"El coeficiente es **positivo** ({coef:.4f}). Respalda la hipótesis marginal de 'engrasar las ruedas' (grease the wheels) donde la opacidad agiliza burocracias."
                else:
                    texto += "La opacidad institucional medida por la corrupción real no muestra un impacto directo y lineal sobre la atracción de IED."
                    
            elif tipo_modelo == "CPI_to_IED":
                if significativo:
                    if coef > 0:
                        texto += f"El coeficiente es **positivo** ({coef:.4f}) y significativo. Un aumento en la transparencia (CPI Score) {t_anterior} una **mayor atracción de inversión extranjera directa**, probando que los inversionistas priorizan la certeza institucional."
                    else:
                        texto += f"El coeficiente es **negativo** ({coef:.4f}). Mayor transparencia reduce los flujos de capital."
                else:
                    texto += "La puntuación del CPI no exhibe un impacto estadísticamente significativo sobre los niveles de IED."
                    
            elif tipo_modelo == "RSF_to_IED":
                if significativo:
                    if coef > 0:
                        texto += f"El coeficiente es **positivo** ({coef:.4f}) y significativo. La libertad de prensa ejerce un beneficio económico directo sobre la IED, reduciendo las asimetrías de información del mercado exterior."
                    else:
                        texto += f"El coeficiente es **negativo** ({coef:.4f})."
                else:
                    texto += "La libertad de prensa no muestra un impacto directo significativo sobre la IED fuera de su canal mediador institucional."
            return texto

        resultados_exportacion, interpretaciones_exportacion = {}, {}

        def correr_modelo_panel(num_mod, titulo, var_dep, vars_indep, var_foco, tipo_clave):
            st.markdown(f"#### {titulo}")
            variables_necesarias = [var_dep] + vars_indep
            df_m = df_panel[variables_necesarias].dropna()
            
            if not df_m.empty:
                X = sm.add_constant(df_m[vars_indep])
                Y = df_m[var_dep]
                mod = PanelOLS(Y, X, entity_effects=True, time_effects=True)
                res = mod.fit(cov_type='robust')
                st.text(res.summary)
                
                interp = interpretar_resultado_avanzado(tipo_clave, var_foco, res.params[var_foco], res.pvalues[var_foco], aplicar_rezago_global)
                st.info("💡 " + interp)
                
                resultados_exportacion[titulo] = res.summary.as_text()
                interpretaciones_exportacion[titulo] = interp
            else:
                st.warning(f"Insuficientes datos válidos para computar el {titulo}.")

        # BLOQUE 1: LIBERTAD DE PRENSA COMO PREDICTOR DE INSTITUCIONES
        st.markdown("### 🏛️ Bloque 1: Impacto de la Libertad de Prensa sobre la Calidad Institucional")
        vars_m1 = ['RSF_L1', 'CORRUPCION_L1'] if aplicar_rezago_global else ['RSF_SCORE']
        correr_modelo_panel("1A", "Modelo 1A: Efecto de Prensa sobre Corrupción Real", 'CORRUPCION', vars_m1, var_rsf_explicativa, "RSF_to_CORR")
        
        vars_m1b = ['RSF_L1', 'CPI_L1'] if aplicar_rezago_global else ['RSF_SCORE']
        correr_modelo_panel("1B", "Modelo 1B: Efecto de Prensa sobre el CPI Score (Transparencia)", 'CPI_SCORE', vars_m1b, var_rsf_explicativa, "RSF_to_CPI")

        # BLOQUE 2: RELACIÓN INVERTIDA (RETROALIMENTACIÓN)
        st.markdown("### 🔄 Bloque 2: Modelos Invertidos (Efecto de las Instituciones sobre la Prensa)")
        vars_m2a = ['CORRUPCION_L1', 'RSF_L1'] if aplicar_rezago_global else ['CORRUPCION']
        correr_modelo_panel("2A", "Modelo 2A Invertido: Efecto de Corrupción Real sobre la Prensa", 'RSF_SCORE', vars_m2a, var_corr_explicativa, "CORR_to_RSF")
        
        vars_m2b = ['CPI_L1', 'RSF_L1'] if aplicar_rezago_global else ['CPI_SCORE']
        correr_modelo_panel("2B", "Modelo 2B Invertido: Efecto del CPI Score sobre la Prensa", 'RSF_SCORE', vars_m2b, var_cpi_explicativa, "CPI_to_RSF")

        # BLOQUE 3: IMPACTO SOBRE LA INVERSIÓN EXTRANJERA DIRECTA (IED)
        if 'IED_PIB' in df_panel.columns:
            st.markdown("### 📈 Bloque 3: Determinantes de la Inversión Extranjera Directa (IED)")
            vars_m3a = ['CORRUPCION_L1', 'IED_L1'] if aplicar_rezago_global else ['CORRUPCION']
            correr_modelo_panel("3A", "Modelo 3A: Efecto de Corrupción Real sobre la IED", 'IED_PIB', vars_m3a, var_corr_explicativa, "CORR_to_IED")
            
            vars_m3b = ['CPI_L1', 'IED_L1'] if aplicar_rezago_global else ['CPI_SCORE']
            correr_modelo_panel("3B", "Modelo 3B: Efecto del CPI Score (Transparencia) sobre la IED", 'IED_PIB', vars_m3b, var_cpi_explicativa, "CPI_to_IED")
            
            vars_m4 = ['RSF_L1', 'IED_L1'] if aplicar_rezago_global else ['RSF_SCORE']
            correr_modelo_panel("4", "Modelo 4: Efecto Directo de la Libertad de Prensa sobre la IED", 'IED_PIB', vars_m4, var_rsf_explicativa, "RSF_to_IED")

    # =========================================================================
    # PESTAÑA 4: COMPILACIÓN Y EXPORTACIÓN DE REPORTES INSTITUCIONALES
    # =========================================================================
    with tab4:
        st.subheader("Módulo de Descargas e Informes Académicos")
        
        output_excel = io.BytesIO()
        df_long.to_excel(output_excel, index=False, engine='openpyxl')
        st.download_button(
            label="📥 Descargar Base de Datos Transformada (Excel - Formato Largo)",
            data=output_excel.getvalue(),
            file_name="Panel_LATAM_LongFormat.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
        
        def generar_word_reporte():
            doc = docx.Document()
            doc.add_heading('Informe de Investigación Econométrica: Libertad de Prensa, Corrupción e IED', 0)
            
            tipo_mod = "Dinámicos con Rezagos Completos (t-1)" if aplicar_rezago_global else "Estáticos Contemporáneos de Control (t)"
            doc.add_paragraph(f"Nota Metodológica Integrada: Este reporte compila los resultados utilizando una especificación de modelos {tipo_mod}.")
            
            doc.add_heading('1. Análisis Descriptivo Agregado del Panel', level=1)
            doc.add_paragraph(desc_stats.to_string())
            
            doc.add_heading('2. Estimaciones de Regresión de Panel (Efectos Fijos Bidimensionales)', level=1)
            
            for nombre_modelo, resumen_texto in resultados_exportacion.items():
                doc.add_heading(nombre_modelo, level=2)
                doc.add_paragraph(resumen_texto)
                if nombre_modelo in interpretaciones_exportacion:
                    doc.add_paragraph(interpretaciones_exportacion[nombre_modelo])
            
            word_io = io.BytesIO()
            doc.save(word_io)
            return word_io.getvalue()
            
        word_data = generar_word_reporte()
        st.download_button(
            label="📄 Descargar Informe Completo Ampliado (Word)",
            data=word_data,
            file_name="Reporte_Final_Econometrico_Ampliado.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
